import io
import pathlib
import re

import docx
from . import utils

heading_re = re.compile(r'Heading (\d+)$')


def prepare_plain_text(
    doc_path: str,
    add_empty_lines: bool = False,
    add_tabs: bool = False,
    upper_headers: bool = False,
    headers_spacing: bool = False,
    remove_empty_lines: bool = False,
    tab_size: int = 8,
):
    doc = docx.Document(doc_path)
    out = io.StringIO()

    prev_has_text = True
    for par in doc.paragraphs:
        style_name = par.style.name
        if style_name.startswith('Heading'):
            print(file=out)
            if headers_spacing:
                print(file=out)

            h_text = par.text
            if upper_headers:
                h_text = h_text.upper()
            elif m := heading_re.match(style_name):
                level = int(m.group(1))
                for _ in range(level):
                    print('#', end='', file=out)
                print(' ', end='', file=out)
            else:
                print('# ', end='', file=out)

            print(h_text, file=out)
            if headers_spacing:
                print(file=out)
            continue

        if remove_empty_lines and prev_has_text and (not par.text or par.text.isspace()):
            prev_has_text = False
            continue

        prev_has_text = True

        if add_tabs:
            print('\xA0' * tab_size, end='', file=out)

        print(par.text, file=out)
        if add_empty_lines:
            print(file=out)

    print(out.getvalue().rstrip())


def partition(doc_path: str, out_dir: str):
    out_dir = pathlib.Path(out_dir).expanduser()
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = docx.Document(doc_path)

    part_num = 0
    first_part = True

    out_doc = docx.Document()

    def init_out_doc():
        nonlocal out_doc
        out_doc = docx.Document(doc_path)
        utils.clear_doc(out_doc)

    init_out_doc()

    def save_doc():
        nonlocal part_num, first_part
        if first_part:
            first_part = False
            return

        part_num += 1
        out_file = out_dir / f'part-{part_num:04d}.docx'
        out_doc.save(str(out_file))
        init_out_doc()

    for par in doc.paragraphs:
        if par.style.name in ('Heading', 'Heading 1'):
            save_doc()
        utils.copy_paragraph(out_doc, par)

    if out_doc.paragraphs:
        save_doc()
