from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def copy_paragraph(doc: Document, paragraph: Paragraph):
    output_paragraph = doc.add_paragraph(style=paragraph.style)

    # Alignment data of whole paragraph
    output_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
    for row in paragraph.runs:
        output_row: Run = output_paragraph.add_run(row.text, style=row.style)
        # Font data
        output_row.style.name = row.style.name
        # Size of font data
        output_row.font.size = row.font.size
        # Bold data
        output_row.bold = row.bold
        # Italic data
        output_row.italic = row.italic
        # Underline data
        output_row.underline = row.underline
        # Color data
        output_row.font.color.rgb = row.font.color.rgb


def clear_doc(doc: Document):
    for par in doc.paragraphs:
        # noinspection PyProtectedMember
        p = par._element
        p.getparent().remove(p)
        p._p = p._element = None
